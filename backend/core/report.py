# -*- coding: utf-8 -*-
"""Word (.docx) taşkın raporu üreteci — Boztepe Bölüm 4.7.x biçiminde.

Hesap sonuçlarından (DSİ Sentetik + opsiyonel Snyder/Mockus/Rasyonel) anlatım
metni, tablolar (yinelenme pikleri, hidrograf koordinatları) ve şekiller
(matplotlib ile çizilen taşkın hidrografları) içeren bir bölüm oluşturur.
"""
import io

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt  # noqa: E402
from docx import Document  # noqa: E402
from docx.enum.table import WD_TABLE_ALIGNMENT  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.shared import Cm, Pt, RGBColor  # noqa: E402

DURS = [2, 4, 6, 8, 12, 18, 24]
RP_ALL = ["2", "5", "10", "25", "50", "100", "500", "1000", "10000", "OET"]
RP_HYD = ["2", "5", "10", "25", "50", "100", "OET"]
DT = 0.5
TEAL = RGBColor(0x0D, 0x5C, 0x63)


def _n(x, d=2):
    if x is None:
        return "—"
    return f"{x:,.{d}f}".replace(",", "@").replace(".", ",").replace("@", ".")


def _governing_dur(sonuc, rp):
    best, pk = DURS[0], -1
    for d in DURS:
        v = sonuc["kabulet"].get(str(d), {}).get(rp)
        if v is not None and v > pk:
            pk, best = v, d
    return best, pk


# ------------------------------------------------------------------ şekiller
def _fig_hydrographs(sonuc, rps, title):
    """Verilen tekerrürlerin taşkın hidrograflarını tek grafikte çizer -> PNG."""
    colors = {"2": "#9db5b2", "5": "#64b5aa", "10": "#2a9d8f", "25": "#d9a441",
              "50": "#e07b3a", "100": "#c73e3a", "OET": "#5e2d48"}
    fig, ax = plt.subplots(figsize=(6.4, 3.6), dpi=150)
    for rp in rps:
        d, _ = _governing_dur(sonuc, rp)
        arr = sonuc["dsi"]["hidrograflar"].get(str(d), {}).get(rp)
        if not arr:
            continue
        t = [i * DT for i in range(len(arr))]
        ax.plot(t, arr, label=f"Q{rp}", color=colors.get(rp, "#333"), linewidth=1.4)
    ax.set_xlabel("Zaman (saat)")
    ax.set_ylabel("Debi (m³/s)")
    ax.set_title(title, fontsize=9)
    ax.grid(True, alpha=0.3)
    ax.legend(fontsize=7, ncol=4)
    fig.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png")
    plt.close(fig)
    buf.seek(0)
    return buf


def _fig_single(sonuc, rp, title):
    d, _ = _governing_dur(sonuc, rp)
    arr = sonuc["dsi"]["hidrograflar"].get(str(d), {}).get(rp)
    fig, ax = plt.subplots(figsize=(6.4, 3.4), dpi=150)
    if arr:
        t = [i * DT for i in range(len(arr))]
        ax.fill_between(t, arr, color="#2a9d8f", alpha=0.25)
        ax.plot(t, arr, color="#0d5c63", linewidth=1.6)
    ax.set_xlabel("Zaman (saat)")
    ax.set_ylabel("Debi (m³/s)")
    ax.set_title(title, fontsize=9)
    ax.grid(True, alpha=0.3)
    fig.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png")
    plt.close(fig)
    buf.seek(0)
    return buf


# ------------------------------------------------------------------ docx yard.
def _heading(doc, text, level, num=None):
    p = doc.add_paragraph()
    run = p.add_run((num + " " if num else "") + text)
    run.bold = True
    run.font.size = Pt(13 - level)
    run.font.color.rgb = TEAL
    p.space_after = Pt(4)
    return p


def _caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)


def _table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, hcell in enumerate(t.rows[0].cells):
        hcell.text = str(headers[j])
        for para in hcell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in para.runs:
                r.bold = True
                r.font.size = Pt(8)
    for row in rows:
        cells = t.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = str(val)
            for para in cells[j].paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER if j else WD_ALIGN_PARAGRAPH.LEFT
                for r in para.runs:
                    r.font.size = Pt(8)
    return t


# ------------------------------------------------------------------ ana üretici
def build_report(girdi, sonuc, meta=None):
    meta = meta or {}
    ad = meta.get("proje_adi") or girdi.get("ad") or sonuc["girdi_ozeti"].get("ad") or "Proje"
    bolum = meta.get("bolum_no", "4.7.3")
    mf = meta.get("MF", 1.13)
    thiessen = meta.get("thiessen") or []
    go = sonuc["girdi_ozeti"]
    on = sonuc["dsi_onhesap"]

    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(10.5)

    _heading(doc, f"{ad} — Sentetik Yöntemlerle Taşkın Tekerrür Debileri Hesabı",
             1, bolum)

    # --- Baz Akım
    _heading(doc, "Baz Akım", 2, bolum + ".1")
    qbaz = go.get("Qbaz") or 0
    if qbaz > 0:
        doc.add_paragraph(
            f"Proje havzasının baz akımı {_n(qbaz)} m³/s olarak alınmış ve tüm "
            f"yinelenmeli taşkın hidrograflarına eklenmiştir.")
    else:
        doc.add_paragraph("Proje havzasında baz akım ihmal edilmiştir.")

    # --- Yağış Analizi
    _heading(doc, "Yağış Analizi", 2, bolum + ".2")
    thi_txt = ""
    if thiessen:
        pay = ", ".join(f"{t['name']} (%{_n(t['agirlik'] * 100, 0)})"
                        for t in thiessen if t.get("agirlik", 0) > 0)
        thi_txt = (f" Yağış alanını temsil eden meteoroloji istasyonları arasında "
                   f"Thiessen poligonu çizilmiş ve temsil oranları hesaplanmıştır: {pay}.")
    doc.add_paragraph(
        "Proje alanı kritik yağış süresindeki yağışları; plüviyograf oranları (PLV), "
        f"alan dağılım katsayısı (ADK/YAD), maksimize faktörü (MF={_n(mf)}) ve Thiessen "
        "oranı ile çarpılarak hesaplanmıştır." + thi_txt +
        " Havzayı temsil eden istasyonların ağırlıklı 24 saatlik yinelenmeli yağış "
        "değerleri aşağıdaki tabloda verilmiştir.")
    p24 = girdi.get("P24", {})
    rows = [["Ağırlıklı yağış (mm)"] + [_n(p24.get(str(t), p24.get(t))) for t in [2, 5, 10, 25, 50, 100]]
            + [_n(girdi.get("P24_OET"))]]
    _table(doc, ["Yinelenme", "2", "5", "10", "25", "50", "100", "OEY"], rows)
    _caption(doc, "Tablo 1  Ağırlıklı 24 Saatlik Yinelenmeli Yağışlar (mm)")

    # --- yöntem seçimi
    ctr = {"t": 1, "s": 0}  # tablo/şekil sayaçları (yağış tablosu = Tablo 1)
    order = ["dsi", "mockus", "rasyonel", "snyder"]
    present = [m for m in order if _method_present(sonuc, m)]
    dahil = [m for m in (meta.get("rapor_yontemleri") or present) if m in present]
    if not dahil:
        dahil = present
    secili = meta.get("secili_yontem")
    if secili not in dahil:
        secili = dahil[0] if dahil else None

    sec = 3  # bölüm.3'ten itibaren yöntem bölümleri
    for m in dahil:
        num = f"{bolum}.{sec}"
        if m == "dsi":
            _dsi_section(doc, sonuc, go, on, num, ad, ctr)
        elif m == "mockus":
            _mockus_section(doc, sonuc, num, ad, ctr)
        elif m == "rasyonel":
            _rasyonel_section(doc, sonuc, num, ad, ctr)
        elif m == "snyder":
            _snyder_section(doc, sonuc["snyder"], num, ad, ctr)
        sec += 1

    # --- karşılaştırma (>=2 dahil yöntem)
    if len(dahil) >= 2:
        _comparison_section(doc, sonuc, f"{bolum}.{sec}", dahil, secili, ctr)
        sec += 1

    # --- seçilen (kabul edilen) yöntem — sonuç
    if secili:
        _conclusion_section(doc, sonuc, f"{bolum}.{sec}", dahil, secili, ctr)

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.getvalue()


# ----------------------------------------------------------- yöntem yardımcıları
METHOD_NAMES = {"dsi": "DSİ Sentetik Yöntemi", "mockus": "Mockus Yöntemi",
                "rasyonel": "Rasyonel Yöntem", "snyder": "Snyder Yöntemi"}


def _method_present(sonuc, m):
    if m == "dsi":
        return bool(sonuc.get("kabulet"))
    return bool(sonuc.get(m))


def _peaks_of(sonuc, m):
    """Yöntem için {rp: pik} (m³/s). DSİ=süre zarfı, Mockus=K1, Rasyonel OET yok."""
    out = {}
    if m == "dsi":
        for rp in RP_ALL:
            vals = [sonuc["kabulet"].get(str(d), {}).get(rp) for d in DURS]
            out[rp] = max([v for v in vals if v is not None], default=None)
    elif m == "mockus":
        s = sonuc["mockus"]["sonuclar"]["K1"]
        for rp in RP_ALL:
            out[rp] = (s["Q_OET"] if rp == "OET" else
                       s["Q_ext"].get(rp) if rp in ("500", "1000", "10000") else s["Q"].get(rp))
    elif m == "rasyonel":
        r = sonuc["rasyonel"]
        for rp in RP_ALL:
            out[rp] = (None if rp == "OET" else
                       r["Q_ext"].get(rp) if rp in ("500", "1000", "10000") else r["Q"].get(rp))
    elif m == "snyder":
        for rp in RP_ALL:
            out[rp] = sonuc["snyder"]["pikler"].get(rp)
    return out


def _peak_table(doc, ctr, sonuc, m, name):
    ctr["t"] += 1
    pk = _peaks_of(sonuc, m)
    rows = [[f"{name} (m³/s)"] + [_n(pk.get(rp)) for rp in RP_ALL]]
    _table(doc, ["Yinelenme"] + [f"Q{rp}" for rp in RP_ALL], rows)
    _caption(doc, f"Tablo {ctr['t']}  {name} ile Hesaplanan Taşkın Yinelenme Pikleri (m³/s)")


# ---------------------------------------------------------------- bölümler
def _dsi_section(doc, sonuc, go, on, num, ad, ctr):
    _heading(doc, "DSİ Sentetik Yöntemi ile Taşkın Hesabı", 2, num)
    doc.add_paragraph(
        f"{ad} havzası; yağış alanı A={_n(go['A_km2'], 3)} km², en uzun akarsu kol boyu "
        f"L={_n(go['L_km'], 3)} km, yağış alanı ağırlık merkezinin en uzun akarsu kol boyu "
        f"üzerine izdüşüm noktası ile proje kesiti arasındaki mesafe Lc={_n(go['Lc_km'], 3)} km "
        f"olarak ölçülmüştür. Harmonik eğim S={_n(go['S_harmonik'], 5)}, yağış-akış eğri "
        f"numarası CNII={_n(go['CN2'], 0)} (CNIII={_n(go['CN3'], 0)}), YZD bölge sınıfı "
        f"“{go.get('bolge')}” alınmıştır.")
    doc.add_paragraph(
        f"DSİ-Sentetik yöntemi ile CNII={_n(go['CN2'], 0)} no'lu eğriden yinelenmeli yağış "
        "değerleri akışa geçirilip artım akış değerleri bulunarak birim hidrografta 2'şer saat "
        "ötelenerek 2, 4, 6, 8, 12, 18, 24 saat süreli yağmurdan oluşan taşkın yinelenme "
        "değerleri, baz akım eklenerek hesaplanmıştır. Yapılan çalışmada birim hidrograf pik "
        f"debisi Qp={_n(on['Qp'], 4)} m³/s/mm, pike varış süresi Tp={_n(on['Tp'])} saat, taban "
        f"süresi T={_n(on['T_saat'], 0)} saat, birim pik qp={_n(on['qp'])} l/s/km²/mm "
        "hesaplanmıştır. Farklı sağanak sürelerine ait taşkın yinelenme pikleri aşağıdaki "
        "tabloda verilmiştir.")
    ctr["t"] += 1
    hdr = ["Yinelenme (yıl)"] + [f"{d} saat" for d in DURS] + ["Zarf (maks)"]
    rows = []
    for rp in RP_ALL:
        vals = [sonuc["kabulet"].get(str(d), {}).get(rp) for d in DURS]
        mx = max([v for v in vals if v is not None], default=None)
        rows.append([f"Q{rp}"] + [_n(v) for v in vals] + [_n(mx)])
    _table(doc, hdr, rows)
    _caption(doc, f"Tablo {ctr['t']}  DSİ Sentetik Yöntemi ile Hesaplanan Taşkın Yinelenme Pikleri (m³/s)")

    # hidrograf koordinatları
    govs = {rp: _governing_dur(sonuc, rp)[0] for rp in RP_HYD}
    maxlen = max(len(sonuc["dsi"]["hidrograflar"][str(govs[rp])][rp]) for rp in RP_HYD)
    ctr["t"] += 1
    hdr = ["T (saat)"] + [f"Q{rp}" for rp in RP_HYD]
    rows = []
    step = 2 if maxlen > 60 else 1
    for i in range(0, maxlen, step):
        row = [_n(i * DT, 1)]
        for rp in RP_HYD:
            arr = sonuc["dsi"]["hidrograflar"][str(govs[rp])][rp]
            row.append(_n(arr[i]) if i < len(arr) else "—")
        rows.append(row)
    _table(doc, hdr, rows)
    dur_note = ", ".join(f"Q{rp}={govs[rp]}sa" for rp in RP_HYD)
    _caption(doc, f"Tablo {ctr['t']}  Yinelenmeli Taşkın Hidrografı Koordinatları (m³/s) — hakim süreler: {dur_note}")

    doc.add_paragraph()
    ctr["s"] += 1
    doc.add_picture(_fig_hydrographs(sonuc, RP_HYD, f"{ad} Taşkın Hidrografları"), width=Cm(15))
    _caption(doc, f"Şekil {ctr['s']}  DSİ Sentetik Yinelenmeli Taşkın Hidrografları (tümü)")
    for rp in [r for r in ["2", "10", "100", "OET"] if r in RP_HYD]:
        yil = "OEY" if rp == "OET" else f"{rp} Yıllık"
        ctr["s"] += 1
        doc.add_picture(_fig_single(sonuc, rp, f"{ad} {yil} Taşkın Hidrografı"), width=Cm(14))
        _caption(doc, f"Şekil {ctr['s']}  {ad} {yil} Taşkın Hidrografı")


def _snyder_section(doc, sn, num, ad, ctr):
    p = sn["parametreler"]
    _heading(doc, "Snyder Yöntemi ile Taşkın Hesabı", 2, num)
    nb = round(24 / p["tr"]) if p.get("tr") else "?"
    doc.add_paragraph(
        f"Snyder yönteminde havza gecikmesi tp=Ct·(L·Lc)^0,30={_n(p['tp'])} saat, standart "
        f"yağış süresi tr={_n(p['tr'], 0)} saat, birim pik qp=2760·Cp/tp={_n(p['qp'])} l/s/km²/cm, "
        f"pik debi Qp={_n(p['Qp'], 3)} m³/s/mm, pike varış Tp={_n(p['Tp'], 0)} saat ve taban "
        f"süresi Tb={_n(p['Tb'], 0)} saat hesaplanmıştır (Ct, Cp katsayıları ile W50={_n(p['W50'], 1)}, "
        f"W75={_n(p['W75'], 1)} saat). 24 saatlik tasarım sağanağı {nb} adet {_n(p['tr'], 0)} saatlik "
        f"bloğa ayrılarak (YZDO dağılımı, ADK={_n(p['YALD'], 3)}, MF=1,13, SCS akış) süperpoze "
        "edilmiş, baz akım eklenmiştir. Taşkın yinelenme pikleri aşağıda verilmiştir.")
    _peak_table(doc, ctr, {"snyder": sn}, "snyder", "Snyder Yöntemi")


def _mockus_section(doc, sonuc, num, ad, ctr):
    m = sonuc["mockus"]
    _heading(doc, "Mockus Yöntemi ile Taşkın Hesabı", 2, num)
    doc.add_paragraph(
        f"Mockus (süperpozesiz) yöntemi ile toplanma süresi Tc={_n(m['Tc'], 3)} saat, yağış "
        f"süresi D={_n(m['D'])} saat, pike varış süresi Tp={_n(m['Tp'], 3)} saat hesaplanmıştır. "
        "Üç akım katsayısı (K1=0,208, K2=0,163, K3 havza özelliğine bağlı) için pik debiler "
        "belirlenmiştir. Raporda temsilci olarak K1 katsayısı sonuçları kullanılmıştır. "
        "Taşkın yinelenme pikleri aşağıda verilmiştir.")
    _peak_table(doc, ctr, sonuc, "mockus", "Mockus Yöntemi (K1)")


def _rasyonel_section(doc, sonuc, num, ad, ctr):
    r = sonuc["rasyonel"]
    _heading(doc, "Rasyonel Yöntem ile Taşkın Hesabı", 2, num)
    doc.add_paragraph(
        f"Rasyonel yöntem (A ≤ 1 km² havzalar için) ile toplanma süresi Tc={_n(r['Tc_dk'], 1)} "
        f"dakika, doğrusal eğim S={_n(r['S_dogrusal'], 5)}, alansal azaltma YADK={_n(r['YADK'], 3)}, "
        f"akış katsayısı C100={_n(r['C100'], 2)} (üs={_n(r['us'], 2)}) kullanılmıştır. "
        f"Q = YADK·C·I·A/3,6 bağıntısı ile taşkın yinelenme pikleri hesaplanmıştır.")
    _peak_table(doc, ctr, sonuc, "rasyonel", "Rasyonel Yöntem")


def _comparison_section(doc, sonuc, num, dahil, secili, ctr):
    _heading(doc, "Yöntemlerin Karşılaştırılması", 2, num)
    doc.add_paragraph(
        "Uygulanan yöntemlerle hesaplanan taşkın yinelenme pikleri aşağıdaki tabloda "
        "karşılaştırılmıştır. Seçilen (kabul edilen) yöntem koyu olarak gösterilmiştir.")
    ctr["t"] += 1
    hdr = ["Yöntem"] + [f"Q{rp}" for rp in RP_ALL]
    rows, bold_rows = [], []
    for i, m in enumerate(dahil):
        pk = _peaks_of(sonuc, m)
        nm = METHOD_NAMES[m] + (" ✔" if m == secili else "")
        rows.append([nm] + [_n(pk.get(rp)) for rp in RP_ALL])
        if m == secili:
            bold_rows.append(i + 1)  # +1: başlık satırı
    t = _table(doc, hdr, rows)
    for ri in bold_rows:
        for cell in t.rows[ri].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
    _caption(doc, f"Tablo {ctr['t']}  Yöntemlerin Taşkın Yinelenme Piklerinin Karşılaştırması (m³/s)")


def _conclusion_section(doc, sonuc, num, dahil, secili, ctr):
    _heading(doc, "Sonuç ve Seçilen Yöntem", 2, num)
    isimler = ", ".join(METHOD_NAMES[m] for m in dahil)
    doc.add_paragraph(
        f"Proje alanı taşkın yinelenme debileri {isimler} kullanılarak hesaplanmıştır. "
        "Hesaplanan değerler karşılaştırılmış; proje sahasının hidrolojik özelliklerini temsil "
        f"açısından ve DSİ'nin ilgili genelgeleri doğrultusunda, {METHOD_NAMES[secili]} ile "
        "hesaplanan taşkın yinelenme değerlerinin projelendirmede esas alınması uygun "
        "bulunmuştur. Seçilen yönteme göre proje kesiti tasarım taşkın debileri aşağıda "
        "verilmiştir.")
    ctr["t"] += 1
    pk = _peaks_of(sonuc, secili)
    _table(doc, ["Tekerrür (yıl)"] + [f"Q{rp}" for rp in RP_ALL],
           [["Tasarım debisi (m³/s)"] + [_n(pk.get(rp)) for rp in RP_ALL]])
    _caption(doc, f"Tablo {ctr['t']}  Seçilen Yönteme ({METHOD_NAMES[secili]}) Göre Tasarım Taşkın Debileri (m³/s)")

# -*- coding: utf-8 -*-
"""MGM veri talebi dilekçesi (Word .docx / PDF) üreteci.

Örnek: `Ayvalık/Balıkesir MGM Dilekçe` biçimi — METEOROLOJİ GENEL
MÜDÜRLÜĞÜNE başlıklı, istasyon(lar) ve ölçüm dönemi (yıl) belirtilen,
istenen veri türleri sıralanan, altta e-posta/GSM/fatura adresi ve
imza/kaşe (görsel) bloğu olan resmi dilekçe.
"""
import os

_HERE = os.path.dirname(os.path.abspath(__file__))
_ROOT = os.path.abspath(os.path.join(_HERE, "..", ".."))
_ASSET = os.path.join(_ROOT, "data", "dilekce")
DEFAULT_IMZA = os.path.join(_ASSET, "imza_kase_default.png")
FONT_REG = os.path.join(_ASSET, "fonts", "DejaVuSerif.ttf")
FONT_BOLD = os.path.join(_ASSET, "fonts", "DejaVuSerif-Bold.ttf")

DEFAULT_VERI = [
    "Aylık toplam yağışlar (mm)",
    "Aylık maksimum yağışlar (mm)",
    "Standart sürede gözlenen en büyük yağışlar (plv) (mm)",
]

# Balıkesir örneğinden çekilen varsayılan iletişim bilgileri (sol alt blok)
DEFAULTS = {
    "eposta": "ygtkumas@gmail.com",
    "gsm": "+90 505 422 9014",
    "adres": ("M. Kemal Mah. 2118. Cad.\n"
              "46-C Blok No:140 Çankaya-Ankara\n"
              "Vergi No: 0012575819\n"
              "Vergi Dairesi: Maltepe"),
    "imza": "",
    "kase": "",
}


def default_imza_bytes():
    if os.path.exists(DEFAULT_IMZA):
        with open(DEFAULT_IMZA, "rb") as f:
            return f.read()
    return None


def _istasyon_ifadesi(no, ad):
    parts = [str(no).strip(), str(ad).strip()]
    return " ".join(p for p in parts if p) + " Meteoroloji Gözlem İstasyonu"


def _sure_ifadesi(aralik):
    aralik = (aralik or "").strip()
    return (f"{aralik} yıllarını kapsayan gözlem süresine ilişkin" if aralik
            else "istasyonun açılışından itibaren tüm gözlem süresini kapsayan")


def _content(d):
    """Dilekçe metnini yapılandırılmış parça listesi olarak döndürür.

    [(tur, metin)] — tur: 'title' | 'body' | 'item' | 'space'.
    """
    il = (d.get("il") or "").strip()
    sts = [s for s in (d.get("istasyonlar") or [])
           if (str(s.get("ad", "")).strip() or str(s.get("no", "")).strip())]
    veri = [str(v).strip() for v in (d.get("veri_turleri") or DEFAULT_VERI) if str(v).strip()]

    parts = [("title", "METEOROLOJİ GENEL MÜDÜRLÜĞÜNE"), ("space", "")]
    ongiris = (f"{il} ili sınırlarında yürütülmekte olan " if il else "") + \
        "hidrolojik analiz çalışmalarında kullanılmak üzere, "
    if len(sts) <= 1:
        s = sts[0] if sts else {"no": "", "ad": "", "aralik": ""}
        parts.append(("body", ongiris + _istasyon_ifadesi(s.get("no"), s.get("ad")) +
                      "na ait " + _sure_ifadesi(s.get("aralik")) +
                      " aşağıdaki meteorolojik verilere ihtiyaç duyulmaktadır."))
    else:
        parts.append(("body", ongiris + "aşağıda belirtilen Meteoroloji Gözlem "
                      "İstasyonlarına ait, karşılarında gösterilen ölçüm dönemlerini "
                      "kapsayan aşağıdaki meteorolojik verilere ihtiyaç duyulmaktadır."))
        for s in sts:
            ar = (s.get("aralik") or "").strip()
            line = _istasyon_ifadesi(s.get("no"), s.get("ad"))
            line += (f" — {ar} yılları arası" if ar
                     else " — açılışından itibaren tüm gözlem süresi")
            parts.append(("item", line))
    parts.append(("body", "Bu kapsamda, ilgili istasyon" + ("lar" if len(sts) > 1 else "") + "a ait;"))
    for v in veri:
        parts.append(("item", v))
    parts.append(("body", "verilerinin elektronik ortamda paylaşılması hususunda;"))
    parts.append(("body", "Gereğini arz ederim."))
    return parts


def build_docx(d: dict, imza_bytes=None) -> bytes:
    import io
    from docx import Document
    from docx.shared import Pt, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH as AL

    doc = Document()
    sec = doc.sections[0]
    for m in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(sec, m, Cm(2.5))
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    def para(text="", align=None, bold=False):
        p = doc.add_paragraph()
        if align is not None:
            p.alignment = align
        if text:
            p.add_run(text).bold = bold
        return p

    for _ in range(4):
        para("", AL.CENTER)
    para("")
    for tur, txt in _content(d):
        if tur == "title":
            para(txt, AL.CENTER, bold=True)
        elif tur == "item":
            p = para(txt)
            p.paragraph_format.left_indent = Cm(1.25)
        elif tur == "space":
            para("")
        else:
            para(txt, AL.JUSTIFY if txt[:3].isalpha() and len(txt) > 60 else None)

    for _ in range(2):
        para("")
    tbl = doc.add_table(rows=1, cols=2)
    left, right = tbl.cell(0, 0), tbl.cell(0, 1)
    left.paragraphs[0].add_run("E-posta: " + (d.get("eposta") or ""))
    left.add_paragraph("GSM: " + (d.get("gsm") or ""))
    left.add_paragraph("Fatura Adresi:")
    for ln in (d.get("adres") or "").split("\n"):
        if ln.strip():
            left.add_paragraph(ln.strip())
    rp = right.paragraphs[0]
    rp.alignment = AL.CENTER
    if imza_bytes:
        rp.add_run().add_picture(io.BytesIO(imza_bytes), width=Inches(2.4))
    else:
        rp.add_run("İmza / Kaşe")
    if (d.get("imza") or "").strip():
        right.add_paragraph((d.get("imza")).strip()).alignment = AL.CENTER
    if (d.get("kase") or "").strip():
        right.add_paragraph((d.get("kase")).strip()).alignment = AL.CENTER

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_pdf(d: dict, imza_bytes=None) -> bytes:
    import io
    from fpdf import FPDF
    from PIL import Image

    pdf = FPDF(format="A4", unit="mm")
    pdf.set_margins(25, 25, 25)
    pdf.set_auto_page_break(True, 25)
    pdf.add_page()
    pdf.add_font("serif", "", FONT_REG)
    pdf.add_font("serif", "B", FONT_BOLD)
    pdf.ln(32)                                   # antet/başlık boşluğu

    for tur, txt in _content(d):
        if tur == "title":
            pdf.set_font("serif", "B", 12)
            pdf.cell(0, 8, txt, align="C", new_x="LMARGIN", new_y="NEXT")
            pdf.ln(3)
        elif tur == "space":
            pdf.ln(3)
        elif tur == "item":
            pdf.set_font("serif", "", 12)
            x = pdf.get_x()
            pdf.set_x(x + 10)
            pdf.multi_cell(0, 6, "•  " + txt, align="L", new_x="LMARGIN", new_y="NEXT")
        else:
            pdf.set_font("serif", "", 12)
            pdf.multi_cell(0, 6, txt, align="J" if len(txt) > 60 else "L",
                           new_x="LMARGIN", new_y="NEXT")
            pdf.ln(1)

    # imza/iletişim bloğu
    y0 = pdf.get_y() + 14
    pdf.set_font("serif", "", 12)
    pdf.set_xy(25, y0)
    adres_lines = "\n".join(ln.strip() for ln in (d.get("adres") or "").split("\n") if ln.strip())
    left_txt = f"E-posta: {d.get('eposta') or ''}\nGSM: {d.get('gsm') or ''}\nFatura Adresi:\n{adres_lines}"
    pdf.multi_cell(85, 6, left_txt, align="L")
    if imza_bytes:
        img = Image.open(io.BytesIO(imza_bytes))
        w = 62.0
        pdf.image(img, x=123, y=y0 - 2, w=w)
        if (d.get("imza") or "").strip():
            pdf.set_xy(120, y0 + w * img.height / img.width)
            pdf.multi_cell(70, 6, (d.get("imza")).strip(), align="C")
    else:
        pdf.set_xy(120, y0)
        pdf.multi_cell(70, 6, "İmza / Kaşe", align="C")

    out = pdf.output()
    return bytes(out)


def build(d: dict, imza_bytes=None, fmt: str = "docx") -> bytes:
    return build_pdf(d, imza_bytes) if fmt == "pdf" else build_docx(d, imza_bytes)
